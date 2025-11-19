import os
import sys
import tkinter as tk
from tkinter import messagebox, filedialog
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageTk

# --- PATH FIX FOR EXE ---
def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller."""
    try:
        base_path = sys._MEIPASS  # type: ignore
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# --- CONFIG ---
RESTAURANT_NAME = "Πιάτα Ημέρας"
LOGO_PATH = resource_path("images/Laterna.png")
NO_SELECTION_IMAGE = resource_path("images/elv.jpg")  # Your custom popup image

# --- MAIN DISHES ---
main_dishes = [
    ("Μοσχαράκι λεμονάτο", "10€", "Lemon veal stew"),
    ("Μουσακάς", "9€", "Moussaka"),
    ("Γεμιστά", "7€", "Stuffed vegetables"),
    ("Μοσχαράκι κοκκινιστό", "11€", "Beef in tomato sauce"),
    ("Μπιφτέκια μοσχαρίσια φούρνου", "9€", "Baked beef patties"),
    ("Αρακάς με πατάτες", "7€", "Peas with potatoes"),
    ("Γίγαντες", "7€", "Giant baked beans"),
    ("Φασόλια μαυρομάτικα με σπανάκι", "6,5€", "Black-eyed beans with spinach"),
    ("Παστίτιο", "8,5€", "Pastitsio"),
    ("Φακές", "6,5€", "Lentil soup"),
    ("Χοιρινό στη γάστρα", "10€", "Pork casserole"),
    ("Φασολάκια κοκκινιστά", "7€", "Green beans in tomato sauce"),
    ("Μπάμιες στο φούρνο", "7€", "Baked okra"),
    ("Μακαρόνια με κιμά μοσχαρίσιο", "8€", "Spaghetti with minced beef"),
    ("Μπριάμ", "7€", "Briam (Greek roasted vegetables)"),
    ("Κοτόπουλο στο φούρνο με πατάτες", "9€", "Oven-baked chicken with potatoes"),
    ("Σουφλέ", "8€", "Soufflé"),
    ("Aρνί στο φούρνο με πατάτες", "12€", "Oven-baked lamb with potatoes")
]

# --- DESSERTS ---
desserts = [
    ("Τριλέτσε", "5,5€", "Soft sponge cake soaked in rich milk sauce with a caramel topping"),
    ("Αμυγδαλόπιτα", "3,5€", "Traditional almond cake with light texture and roasted almond aroma"),
    ("Ραβανί", "4€", "Classic semolina cake with airy texture and light syrup")
]

# --- FULL DISHES LIST ---
dishes = main_dishes + desserts

# --- POPUP WHEN NO SELECTION ---
def show_no_selection_image():
    """Show a large popup image when no dishes are selected."""
    popup = tk.Toplevel()
    popup.title("Προειδοποίηση")
    popup.configure(bg="white")

    screen_w = popup.winfo_screenwidth()
    screen_h = popup.winfo_screenheight()
    win_w = int(screen_w * 0.7)
    win_h = int(screen_h * 0.7)
    pos_x = (screen_w - win_w) // 2
    pos_y = (screen_h - win_h) // 2
    popup.geometry(f"{win_w}x{win_h}+{pos_x}+{pos_y}")
    popup.resizable(False, False)

    try:
        img = Image.open(NO_SELECTION_IMAGE)
        img_ratio = img.width / img.height
        win_ratio = win_w / win_h

        if img_ratio > win_ratio:
            new_w = int(win_w * 0.85)
            new_h = int(new_w / img_ratio)
        else:
            new_h = int(win_h * 0.7)
            new_w = int(new_h * img_ratio)

        img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
        photo = ImageTk.PhotoImage(img)

        label_img = tk.Label(popup, image=photo, bg="white")
        label_img.image = photo
        label_img.pack(pady=20)

    except Exception:
        tk.Label(
            popup,
            text="⚠️ Δεν βρέθηκε η εικόνα!",
            font=("Arial", 16, "bold"),
            fg="red",
            bg="white"
        ).pack(pady=40)

    tk.Label(
        popup,
        text="Επιλέξτε τουλάχιστον ένα πιάτο!",
        font=("Arial", 18, "bold"),
        fg="#D32F2F",
        bg="white"
    ).pack(pady=20)

    tk.Button(
        popup,
        text="OK",
        command=popup.destroy,
        bg="#4CAF50",
        fg="white",
        font=("Arial", 14, "bold"),
        padx=20,
        pady=10
    ).pack(pady=10)

    popup.grab_set()

# --- CREATE WORD DOCUMENT ---
def create_word_document(selected_dishes):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.05)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # --- Logo ---
    try:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.6))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception:
        messagebox.showwarning(
            "Προειδοποίηση",
            f"Δεν βρέθηκε το λογότυπο ({LOGO_PATH}). Θα συνεχιστεί χωρίς αυτό."
        )

    # --- Restaurant name ---
    header = document.add_paragraph(RESTAURANT_NAME)
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header.paragraph_format.space_before = Pt(0)
    header.paragraph_format.space_after = Pt(4)
    run = header.runs[0]
    run.font.bold = True
    run.font.size = Pt(20)

    # --- Separator ---
    document.add_paragraph("──────────────").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Separate Main Dishes and Desserts ---
    selected_main = [d for d in selected_dishes if d in main_dishes]
    selected_desserts = [d for d in selected_dishes if d in desserts]

    if selected_main:
        doc_main_header = document.add_paragraph("Κύρια Πιάτα / Main Dishes")
        doc_main_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc_main_header.runs[0].bold = True
        doc_main_header.runs[0].font.size = Pt(14)
        for dish_name, dish_price, dish_description in selected_main:
            p = document.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            run_greek = p.add_run(f"{dish_name}  ({dish_price})\n")
            run_greek.bold = True
            run_greek.font.size = Pt(13)
            run_english = p.add_run(dish_description)
            run_english.font.size = Pt(10.5)
            run_english.italic = True

    if selected_desserts:
        doc_dessert_header = document.add_paragraph("Γλυκά / Desserts")
        doc_dessert_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc_dessert_header.runs[0].bold = True
        doc_dessert_header.runs[0].font.size = Pt(14)
        for dish_name, dish_price, dish_description in selected_desserts:
            p = document.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            run_greek = p.add_run(f"{dish_name}  ({dish_price})\n")
            run_greek.bold = True
            run_greek.font.size = Pt(13)
            run_english = p.add_run(dish_description)
            run_english.font.size = Pt(10.5)
            run_english.italic = True

    # --- Footer ---
    document.add_paragraph("──────────────").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer = document.add_paragraph("Σας ευχαριστούμε! / Thank you!")
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer.paragraph_format.space_before = Pt(6)

    # --- Save file ---
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")],
        initialfile="Laterna_Menu.docx"
    )
    if not file_path:
        return

    document.save(file_path)
    messagebox.showinfo("Επιτυχία", f"Το αρχείο αποθηκεύτηκε ως:\n{file_path}")

    try:
        os.startfile(file_path)
    except Exception:
        pass

# --- BUTTON HANDLER ---
def on_submit():
    selected_dishes = [dishes[i] for i, var in enumerate(dish_vars) if var.get()]
    if not selected_dishes:
        show_no_selection_image()
        return
    create_word_document(selected_dishes)

# --- GUI SETUP ---
root = tk.Tk()
root.title("Laterna - Dishes of the Day")

tk.Label(root, text="Laterna - Ημερήσια Πιάτα", font=("Arial", 14, "bold")).pack(pady=10)

frame = tk.Frame(root)
frame.pack(fill="both", expand=True, pady=10)

canvas = tk.Canvas(frame, height=350)
scrollbar = tk.Scrollbar(frame, orient="vertical", command=canvas.yview)
scrollable_frame = tk.Frame(canvas)

scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
canvas.configure(yscrollcommand=scrollbar.set)

canvas.pack(side="left", fill="both", expand=True)
scrollbar.pack(side="right", fill="y")

# --- Main Dishes ---
tk.Label(scrollable_frame, text="Επιλέξτε Κύρια Πιάτα:", font=("Arial", 10, "bold")).pack(pady=5)
dish_vars = []
for dish in main_dishes:
    var = tk.BooleanVar()
    chk = tk.Checkbutton(scrollable_frame, text=f"{dish[0]} ({dish[1]})", variable=var, anchor="w")
    chk.pack(anchor="w")
    dish_vars.append(var)

# --- Desserts ---
tk.Label(scrollable_frame, text="Επιλέξτε Γλυκά:", font=("Arial", 10, "bold")).pack(pady=10)
for dish in desserts:
    var = tk.BooleanVar()
    chk = tk.Checkbutton(scrollable_frame, text=f"{dish[0]} ({dish[1]})", variable=var, anchor="w")
    chk.pack(anchor="w")
    dish_vars.append(var)

# --- Submit Button ---
tk.Button(
    root,
    text="ΔΗΜΙΟΥΡΓΙΑ WORD",
    command=on_submit,
    bg="#4CAF50",
    fg="white",
    padx=10,
    pady=5,
    font=("Arial", 11, "bold")
).pack(pady=20)

root.mainloop()
